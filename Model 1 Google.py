import streamlit as st
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.memory import ConversationBufferMemory
import io
import PyPDF2
import docx
import hashlib
from PIL import Image, ImageDraw
import pytesseract
import numpy as np

# Set up Streamlit
st.set_page_config(page_title='AiVocator Junior Counsel', page_icon='AIvocator_icon.png', layout="wide")

# Custom CSS for dark theme and other UI modifications
st.markdown("""
    <style>
    .stApp {
        background-color: #1E1E1E;
        color: #FFFFFF;
    }
    .stButton>button {
        color: #FFFFFF;
        background-color: #4CAF50;
        border-radius: 5px;
    }
    .stTextInput>div>div>input {
        color: #FFFFFF;
        background-color: #2E2E2E;
    }
    .stTextArea>div>div>textarea {
        color: #FFFFFF;
        background-color: #2E2E2E;
    }
    </style>
    """, unsafe_allow_html=True)

# Header with logo
col1, col2 = st.columns([1, 5])
with col1:
    st.image('AIvocator_icon.png', width=50)
with col2:
    st.title('AiVocator Junior Counsel ⚖️')

# Initialize session state
if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(return_messages=True)
if 'history' not in st.session_state:
    st.session_state.history = []
if 'strategies' not in st.session_state:
    st.session_state.strategies = {"plaintiff": "", "defendant": "", "judge": "", "game_theorist": ""}
if 'full_case_info' not in st.session_state:
    st.session_state.full_case_info = ""
if 'show_history' not in st.session_state:
    st.session_state.show_history = False
if 'document_store' not in st.session_state:
    st.session_state.document_store = None
if 'cache' not in st.session_state:
    st.session_state.cache = {}
if 'document_previews' not in st.session_state:
    st.session_state.document_previews = []
if 'uploaded_documents' not in st.session_state:
    st.session_state.uploaded_documents = []

# Input for Google API Key
google_api_key = st.text_input("Enter your Google API Key", type="password")

if not google_api_key:
    st.warning("Please enter your Google API key to proceed.")
    st.stop()

# Initialize the ChatGoogleGenerativeAI model
llm = ChatGoogleGenerativeAI(model='gemini-pro', google_api_key=google_api_key, streaming=True)

# Function to stream responses
def stream_response(response):
    response_placeholder = st.empty()
    full_response = ""
    for chunk in response:
        if isinstance(chunk, dict) and 'text' in chunk:
            full_response += chunk['text']
        elif hasattr(chunk, 'content'):
            full_response += chunk.content
        else:
            full_response += str(chunk)
        response_placeholder.markdown(full_response + "▌")
    response_placeholder.markdown(full_response)
    return full_response

# Function to extract text from various file types, including images using OCR
def extract_text_from_file(file):
    def try_decode(data):
        for encoding in ['utf-8', 'latin1', 'iso-8859-1']:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                pass
        return data.decode('utf-8', errors='ignore')  # Final fallback

    if file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file.type.startswith('image'):
        image = Image.open(file)
        text = pytesseract.image_to_string(image)
    else:
        text = try_decode(file.getvalue())
    return text

def generate_thumbnail(file):
    if file.type.startswith('image'):
        return file
    elif file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        if len(pdf_reader.pages) > 0:
            page = pdf_reader.pages[0]
            text = page.extract_text()[:100]  # Extract first 100 characters
            return create_text_thumbnail(text, file.name)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        if doc.paragraphs:
            text = doc.paragraphs[0].text[:100]  # Extract first 100 characters
            return create_text_thumbnail(text, file.name)
    else:
        text = try_decode(file.getvalue())[:100]  # Extract first 100 characters
        return create_text_thumbnail(text, file.name)

def create_text_thumbnail(text, filename):
    img = Image.new('RGB', (200, 200), color='white')
    d = ImageDraw.Draw(img)
    d.text((10, 10), f"{filename}\n\n{text}", fill=(0, 0, 0))
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr = img_byte_arr.getvalue()
    return img_byte_arr

# Document upload section
st.subheader("Upload Documents")
uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True)
if uploaded_files:
    documents = []
    col1, col2, col3, col4 = st.columns(4)
    for i, file in enumerate(uploaded_files):
        text = extract_text_from_file(file)
        documents.append(text)
        thumbnail = generate_thumbnail(file)
        with [col1, col2, col3, col4][i % 4]:
            if isinstance(thumbnail, bytes):
                st.image(thumbnail, caption=file.name, use_column_width=True)
            else:
                st.image(thumbnail, caption=file.name, use_column_width=True)
    if documents:
        st.success(f"{len(uploaded_files)} documents uploaded and processed successfully!")
        st.info("You can now use the Document Analysis section below to query and summarize the uploaded documents.")
    else:
        st.warning("No text could be extracted from the uploaded files.")

# Display all uploaded documents
if uploaded_files:
    for file in uploaded_files:
        if file not in [doc['file'] for doc in st.session_state.uploaded_documents]:
            st.session_state.uploaded_documents.append({
                'file': file,
                'thumbnail': generate_thumbnail(file)
            })

if st.session_state.uploaded_documents:
    st.subheader("Uploaded Documents")
    cols = st.columns(4)
    for i, doc in enumerate(st.session_state.uploaded_documents):
        with cols[i % 4]:
            if isinstance(doc['thumbnail'], bytes):
                st.image(doc['thumbnail'], caption=doc['file'].name, use_column_width=True)
            else:
                st.image(doc['thumbnail'], caption=doc['file'].name, use_column_width=True)
            if st.button(f"Remove {doc['file'].name}", key=f"remove_{i}"):
                st.session_state.uploaded_documents.pop(i)
                st.experimental_rerun()

# Sidebar for history
with st.sidebar:
    st.session_state.show_history = st.checkbox("Show Conversation History", value=st.session_state.show_history)
    
    if st.session_state.show_history:
        st.header("Conversation History")
        for i, message in enumerate(st.session_state.history):
            if st.button(f"{i+1}. {message[:50]}...", key=f"history_{i}"):
                st.markdown(f"**Full message:**\n\n{message}")

# Function to generate document summary
def generate_document_summary(document_texts):
    summary_prompt = """
    Please provide a concise summary of the following document. 
    Focus on the main topics, key points, and any important details:

    {doc}

    Summary:
    """
    summary_chain = LLMChain(llm=llm, prompt=PromptTemplate(template=summary_prompt, input_variables=["doc"]), memory=st.session_state.memory)
    summaries = []
    for text in document_texts:
        chunks = [text[i:i+2000] for i in range(0, len(text), 2000)]  # Chunk the document text to avoid API limits
        for chunk in chunks:
            summary = stream_response(summary_chain.stream({"doc": chunk}))
            summaries.append(summary)
    return "\n\n".join(summaries)

# Document querying and summarization section
if st.session_state.uploaded_documents:
    st.subheader("Document Analysis")
    
    # Document summarization
    if st.button("Generate Document Summary"):
        with st.spinner('Generating document summary...'):
            try:
                document_texts = [extract_text_from_file(doc['file']) for doc in st.session_state.uploaded_documents]
                summary = generate_document_summary(document_texts)
                st.session_state.history.append(f"Document Summary: {summary}")
                st.write(summary)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
    
    # Document querying
    doc_query = st.text_input("Ask a question about the uploaded documents:")
    if doc_query:
        if st.button("Search Documents"):
            question_template = """
            Based on the following question, provide a detailed answer considering the information in the documents:

            Question: {question}

            Please give a comprehensive response.
            """
            question_prompt = PromptTemplate(template=question_template, input_variables=["question"])
            question_chain = LLMChain(llm=llm, prompt=question_prompt, memory=st.session_state.memory)
            
            with st.spinner('Searching documents...'):
                try:
                    answer = stream_response(question_chain.stream({"question": doc_query}))
                    st.session_state.history.append(f"Document Q: {doc_query}\nA: {answer}")
                    st.write(answer)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

# Main content
st.subheader('Enter a business case topic to get legal strategies.')
case_description = st.text_area('Describe your business case:', height=100)

if case_description:
    if st.button("Generate Follow-up Questions"):
        follow_up_template = """
        You are an experienced lawyer assistant. Your role is to gather relevant information 
        about a business case and provide legal strategies. Based on the following business case, 
        ask follow-up questions to obtain necessary details such as precedents, contracts, 
        agreements, emails, and official reports.

        Business case: {case_description}

        Please provide a list of follow-up questions to gather more relevant details for this case.
        """
        follow_up_prompt = PromptTemplate(template=follow_up_template, input_variables=["case_description"])
        follow_up_chain = LLMChain(llm=llm, prompt=follow_up_prompt, memory=st.session_state.memory)
        
        with st.spinner('Generating follow-up questions...'):
            follow_up_questions = stream_response(follow_up_chain.stream({"case_description": case_description}))
            st.session_state.history.append(f"Follow-up Questions: {follow_up_questions}")

    follow_up_answers = st.text_area("Your answers to the follow-up questions:", height=150)

    if st.button("Generate Legal Strategies"):
        # Plaintiff's lawyer template
        plaintiff_template = """
        You are an experienced lawyer representing the plaintiff in a South African court. Your role is to analyze 
        the given business case and provide strong legal strategies to win the case for your client (the plaintiff). 
        Based on the following information, suggest effective legal strategies and arguments.

        {full_case_info}
        
        Please provide a detailed analysis and effective legal strategies for the plaintiff to win this case.
        """
        plaintiff_prompt = PromptTemplate(template=plaintiff_template, input_variables=["full_case_info"])
        plaintiff_chain = LLMChain(llm=llm, prompt=plaintiff_prompt, memory=st.session_state.memory)
        
        # Defendant's lawyer template
        defendant_template = """
        You are an experienced lawyer representing the defendant in a South African court. Your role is to analyze 
        the given business case and provide strong counter-strategies to win the case for your client (the defendant). 
        Based on the following information, suggest effective legal counter-arguments and strategies.

        {full_case_info}
        
        Please provide a detailed analysis and effective legal counter-strategies for the defendant to win this case.
        """
        defendant_prompt = PromptTemplate(template=defendant_template, input_variables=["full_case_info"])
        defendant_chain = LLMChain(llm=llm, prompt=defendant_prompt, memory=st.session_state.memory)
        
        st.session_state.full_case_info = f"""
        Business case: {case_description}
        
        Additional information from follow-up questions:
        {follow_up_answers}
        """
        
        with st.spinner('Generating legal strategies...'):
            st.session_state.strategies["plaintiff"] = stream_response(plaintiff_chain.stream({
                "full_case_info": st.session_state.full_case_info
            }))
            st.session_state.strategies["defendant"] = stream_response(defendant_chain.stream({
                "full_case_info": st.session_state.full_case_info
            }))
            st.session_state.history.append(f"Plaintiff Strategies: {st.session_state.strategies['plaintiff']}")
            st.session_state.history.append(f"Defendant Strategies: {st.session_state.strategies['defendant']}")

    if st.session_state.strategies["plaintiff"] and st.session_state.strategies["defendant"]:
        if st.button("Generate Judge's Ruling"):
            judge_template = """
            You are a judge in a South African court. Based on the following information, including strategies 
            from both the plaintiff and the defendant, make a ruling on this case. Consider South African law, 
            past judgments, and all presented evidence.

            Case Information:
            {full_case_info}

            Plaintiff's Strategies:
            {plaintiff_strategies}

            Defendant's Strategies:
            {defendant_strategies}

            Please provide a detailed ruling, citing relevant laws and past judgments where applicable.
            """
            judge_prompt = PromptTemplate(template=judge_template, input_variables=["full_case_info", "plaintiff_strategies", "defendant_strategies"])
            judge_chain = LLMChain(llm=llm, prompt=judge_prompt)
            
            with st.spinner('Generating judge\'s ruling...'):
                judge_ruling = judge_chain.run({
                    "full_case_info": st.session_state.full_case_info,
                    "plaintiff_strategies": st.session_state.strategies["plaintiff"],
                    "defendant_strategies": st.session_state.strategies["defendant"]
                })
                st.session_state.strategies["judge"] = judge_ruling
                st.session_state.history.append(f"Judge's Ruling: {judge_ruling}")
                st.write(judge_ruling)

        if st.button("Generate Game Theory Analysis"):
            game_theory_template = """
            You are a game theorist specializing in legal strategy. Analyze the strategies provided by both the 
            plaintiff and the defendant, and suggest optimal negotiation strategies for the plaintiff to increase 
            their chances of winning or achieving a favorable settlement.

            Plaintiff's Strategies:
            {plaintiff_strategies}

            Defendant's Strategies:
            {defendant_strategies}

            Please provide a detailed game theory analysis and suggest optimal negotiation strategies for the plaintiff.
            """
            game_theory_prompt = PromptTemplate(template=game_theory_template, input_variables=["plaintiff_strategies", "defendant_strategies"])
            game_theory_chain = LLMChain(llm=llm, prompt=game_theory_prompt)
            
            with st.spinner('Generating game theory analysis...'):
                game_theory_analysis = game_theory_chain.run({
                    "plaintiff_strategies": st.session_state.strategies["plaintiff"],
                    "defendant_strategies": st.session_state.strategies["defendant"]
                })
                st.session_state.strategies["game_theorist"] = game_theory_analysis
                st.session_state.history.append(f"Game Theory Analysis: {game_theory_analysis}")
                st.write(game_theory_analysis)

    # Download button for strategies
    if any(st.session_state.strategies.values()):
        strategies_text = "\n\n".join([f"{key.capitalize()} Strategies:\n{value}" for key, value in st.session_state.strategies.items() if value])
        bytes_io = io.BytesIO(strategies_text.encode())
        st.download_button(
            label="Download Strategies",
            data=bytes_io,
            file_name="legal_strategies.txt",
            mime="text/plain"
        )

additional_question = st.text_input("Ask additional questions:")
if additional_question:
    if st.button("Get Answer"):
        question_template = """
        Based on the previous conversation and the following question, provide a detailed answer:

        Question: {question}

        Please give a comprehensive response considering all the information discussed so far.
        """
        question_prompt = PromptTemplate(template=question_template, input_variables=["question"])
        question_chain = LLMChain(llm=llm, prompt=question_prompt, memory=st.session_state.memory)
        
        with st.spinner('Generating answer...'):
            answer = stream_response(question_chain.stream({"question": additional_question}))
            st.session_state.history.append(f"Q: {additional_question}\nA: {answer}")
