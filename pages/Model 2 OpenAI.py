import streamlit as st
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.memory import ConversationBufferMemory
from langchain_openai import ChatOpenAI
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.schema import Document
import io
import PyPDF2
import docx
import hashlib
from PIL import Image, ImageDraw
import pytesseract
import numpy as np

# Set up Streamlit
st.set_page_config(page_title='AiVocator Junior Counsel', page_icon='AIvocator_icon.png', layout="wide")

# Custom CSS for dark theme and other UI modifications
st.markdown("""
    <style>
    .stApp {
        background-color: #1E1E1E;
        color: #FFFFFF;
    }
    .stButton>button {
        color: #FFFFFF;
        background-color: #4CAF50;
        border-radius: 5px;
    }
    .stTextInput>div>div>input {
        color: #FFFFFF;
        background-color: #2E2E2E;
    }
    .stTextArea>div>div>textarea {
        color: #FFFFFF;
        background-color: #2E2E2E;
    }
    </style>
    """, unsafe_allow_html=True)

# Header with logo
col1, col2 = st.columns([1, 5])
with col1:
    st.image('AIvocator_icon.png', width=50)
with col2:
    st.title('AiVocator Junior Counsel ⚖️')

# Initialize session state
if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(return_messages=True)
if 'history' not in st.session_state:
    st.session_state.history = []
if 'strategies' not in st.session_state:
    st.session_state.strategies = {"plaintiff": "", "defendant": "", "judge": "", "game_theorist": ""}
if 'full_case_info' not in st.session_state:
    st.session_state.full_case_info = ""
if 'show_history' not in st.session_state:
    st.session_state.show_history = False
if 'document_store' not in st.session_state:
    st.session_state.document_store = None
if 'cache' not in st.session_state:
    st.session_state.cache = {}
if 'document_previews' not in st.session_state:
    st.session_state.document_previews = []
if 'uploaded_documents' not in st.session_state:
    st.session_state.uploaded_documents = []

# Input for OpenAI API Key
openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")

if not openai_api_key:
    st.warning("Please enter your OpenAI API key to proceed.")
    st.stop()

@st.cache_resource
def load_llm(api_key):
    return ChatOpenAI(model_name='gpt-4', openai_api_key=api_key)

llm = load_llm(openai_api_key)

# Function to extract text from various file types, including OCR for images
def extract_text_from_file(file):
    if file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file.type.startswith('image'):
        image = Image.open(file)
        text = pytesseract.image_to_string(image)
    else:
        text = file.getvalue().decode("utf-8")
    return text

# Function to generate thumbnail for documents
def generate_thumbnail(file):
    if file.type.startswith('image'):
        return file
    elif file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        if len(pdf_reader.pages) > 0:
            page = pdf_reader.pages[0]
            text = page.extract_text()[:100]  # Extract first 100 characters
            return create_text_thumbnail(text, file.name)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        if doc.paragraphs:
            text = doc.paragraphs[0].text[:100]  # Extract first 100 characters
            return create_text_thumbnail(text, file.name)
    else:
        text = file.getvalue().decode("utf-8")[:100]  # Extract first 100 characters
        return create_text_thumbnail(text, file.name)

def create_text_thumbnail(text, filename):
    img = Image.new('RGB', (200, 200), color='white')
    d = ImageDraw.Draw(img)
    d.text((10, 10), f"{filename}\n\n{text}", fill=(0, 0, 0))
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr = img_byte_arr.getvalue()
    return img_byte_arr

# Function to create document store
def create_document_store(documents):
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    docs = [Document(page_content=text) for text in documents]
    texts = text_splitter.split_documents(docs)
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
    return FAISS.from_documents(texts, embeddings)

# Function to get answer from document store
def get_answer_from_docs(question):
    if st.session_state.document_store is None:
        return "No documents have been uploaded yet."
    
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=st.session_state.document_store.as_retriever())
    return qa_chain.run(question)

# Function to generate cache key
def generate_cache_key(prompt):
    return hashlib.md5(prompt.encode()).hexdigest()

# Function to get response (with caching)
def get_response(chain, prompt):
    cache_key = generate_cache_key(str(prompt))
    if cache_key in st.session_state.cache:
        return st.session_state.cache[cache_key]
    else:
        response = chain.run(prompt)
        st.session_state.cache[cache_key] = response
        return response

# Sidebar for history
with st.sidebar:
    st.session_state.show_history = st.checkbox("Show Conversation History", value=st.session_state.show_history)
    
    if st.session_state.show_history:
        st.header("Conversation History")
        for i, message in enumerate(st.session_state.history):
            if st.button(f"{i+1}. {message[:50]}...", key=f"history_{i}"):
                st.markdown(f"**Full message:**\n\n{message}")

# Document upload section
st.subheader("Upload Documents")
uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True)
if uploaded_files:
    documents = []
    col1, col2, col3, col4 = st.columns(4)
    for i, file in enumerate(uploaded_files):
        text = extract_text_from_file(file)
        documents.append(text)
        thumbnail = generate_thumbnail(file)
        with [col1, col2, col3, col4][i % 4]:
            if isinstance(thumbnail, bytes):
                st.image(thumbnail, caption=file.name, use_column_width=True)
            else:
                st.image(thumbnail, caption=file.name, use_column_width=True)
    if documents:
        st.session_state.document_store = create_document_store(documents)
        st.success(f"{len(uploaded_files)} documents uploaded and processed successfully!")
        st.info("You can now use the Document Analysis section below to query and summarize the uploaded documents.")
    else:
        st.warning("No text could be extracted from the uploaded files.")

# Display document thumbnails
st.subheader("Uploaded Document Previews")
if st.session_state.document_store:
    cols = st.columns(4)
    for i, doc in enumerate(st.session_state.document_previews):
        with cols[i % 4]:
            st.markdown(f"**Document {i+1}:**")
            st.markdown(f"{doc}...")
else:
    st.markdown("No documents uploaded yet.")

# Function to generate document summary
def generate_document_summary():
    if st.session_state.document_store is None:
        return "No documents have been uploaded yet."
    
    summary_prompt = """
    Please provide a concise summary of the following documents. 
    Focus on the main topics, key points, and any important details:

    {docs}
    
    Summary:
    """
    summary_chain = LLMChain(llm=llm, prompt=PromptTemplate(template=summary_prompt, input_variables=["docs"]))
    docs = st.session_state.document_store.similarity_search("", k=5)
    summaries = []
    for doc in docs:
        doc_content = doc.page_content
        if len(doc_content) > 1000:
            chunks = [doc_content[i:i+1000] for i in range(0, len(doc_content), 1000)]
            for chunk in chunks:
                summary = summary_chain.run(chunk)
                summaries.append(summary)
        else:
            summary = summary_chain.run(doc_content)
            summaries.append(summary)
    return "\n\n".join(summaries)

# Document querying and summarization section
if st.session_state.document_store is not None:
    st.subheader("Document Analysis")
    
    # Document summarization
    if st.button("Generate Document Summary"):
        with st.spinner('Generating document summary...'):
            try:
                summary = generate_document_summary()
                st.session_state.history.append(f"Document Summary: {summary}")
                st.write(summary)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
    
    # Document querying
    doc_query = st.text_input("Ask a question about the uploaded documents:")
    if doc_query:
        if st.button("Search Documents"):
            with st.spinner('Searching documents...'):
                try:
                    answer = get_answer_from_docs(doc_query)
                    st.session_state.history.append(f"Document Q: {doc_query}\nA: {answer}")
                    st.write(answer)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

# Main content
st.subheader('Enter a business case topic to get legal strategies.')
case_description = st.text_area('Describe your business case:', height=100)

if case_description:
    if st.button("Generate Follow-up Questions"):
        follow_up_template = """
        You are an experienced lawyer assistant. Your role is to gather relevant information 
        about a business case and provide legal strategies. Based on the following business case, 
        ask follow-up questions to obtain necessary details such as precedents, contracts, 
        agreements, emails, and official reports.

        Business case: {case_description}

        Please provide a list of follow-up questions to gather more relevant details for this case.
        """
        follow_up_prompt = PromptTemplate(template=follow_up_template, input_variables=["case_description"])
        follow_up_chain = LLMChain(llm=llm, prompt=follow_up_prompt, memory=st.session_state.memory)
        
        with st.spinner('Generating follow-up questions...'):
            try:
                follow_up_questions = get_response(follow_up_chain, {"case_description": case_description})
                st.session_state.history.append(f"Follow-up Questions: {follow_up_questions}")
                st.write(follow_up_questions)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    follow_up_answers = st.text_area("Your answers to the follow-up questions:", height=150)

    if st.button("Generate Legal Strategies"):
        # Plaintiff's lawyer template
        plaintiff_template = """
        You are an experienced lawyer representing the plaintiff in a South African court. Your role is to analyze 
        the given business case and provide strong legal strategies to win the case for your client (the plaintiff). 
        Based on the following information, suggest effective legal strategies and arguments.

        {full_case_info}
        
        Please provide a detailed analysis and effective legal strategies for the plaintiff to win this case.
        """
        plaintiff_prompt = PromptTemplate(template=plaintiff_template, input_variables=["full_case_info"])
        plaintiff_chain = LLMChain(llm=llm, prompt=plaintiff_prompt, memory=st.session_state.memory)
        
        # Defendant's lawyer template
        defendant_template = """
        You are an experienced lawyer representing the defendant in a South African court. Your role is to analyze 
        the given business case and provide strong counter-strategies to win the case for your client (the defendant). 
        Based on the following information, suggest effective legal counter-arguments and strategies.

        {full_case_info}
        
        Please provide a detailed analysis and effective legal counter-strategies for the defendant to win this case.
        """
        defendant_prompt = PromptTemplate(template=defendant_template, input_variables=["full_case_info"])
        defendant_chain = LLMChain(llm=llm, prompt=defendant_prompt, memory=st.session_state.memory)
        
        st.session_state.full_case_info = f"""
        Business case: {case_description}
        
        Additional information from follow-up questions:
        {follow_up_answers}
        """
        
        with st.spinner('Generating legal strategies...'):
            try:
                st.session_state.strategies["plaintiff"] = get_response(plaintiff_chain, {
                    "full_case_info": st.session_state.full_case_info
                })
                st.session_state.strategies["defendant"] = get_response(defendant_chain, {
                    "full_case_info": st.session_state.full_case_info
                })
                st.session_state.history.append(f"Plaintiff Strategies: {st.session_state.strategies['plaintiff']}")
                st.session_state.history.append(f"Defendant Strategies: {st.session_state.strategies['defendant']}")
                st.write(st.session_state.strategies["plaintiff"])
                st.write(st.session_state.strategies["defendant"])
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    if st.session_state.strategies["plaintiff"] and st.session_state.strategies["defendant"]:
        if st.button("Generate Judge's Ruling"):
            judge_template = """
            You are a judge in a South African court. Based on the following information, including strategies 
            from both the plaintiff and the defendant, make a ruling on this case. Consider South African law, 
            past judgments, and all presented evidence.

            Case Information:
            {full_case_info}

            Plaintiff's Strategies:
            {plaintiff_strategies}

            Defendant's Strategies:
            {defendant_strategies}

            Please provide a detailed ruling, citing relevant laws and past judgments where applicable.
            """
            judge_prompt = PromptTemplate(template=judge_template, input_variables=["full_case_info", "plaintiff_strategies", "defendant_strategies"])
            judge_chain = LLMChain(llm=llm, prompt=judge_prompt)
            
            with st.spinner('Generating judge\'s ruling...'):
                try:
                    judge_ruling = get_response(judge_chain, {
                        "full_case_info": st.session_state.full_case_info,
                        "plaintiff_strategies": st.session_state.strategies["plaintiff"],
                        "defendant_strategies": st.session_state.strategies["defendant"]
                    })
                    st.session_state.strategies["judge"] = judge_ruling
                    st.session_state.history.append(f"Judge's Ruling: {judge_ruling}")
                    st.write(judge_ruling)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

        if st.button("Generate Game Theory Analysis"):
            game_theory_template = """
            You are a game theorist specializing in legal strategy. Analyze the strategies provided by both the 
            plaintiff and the defendant, and suggest optimal negotiation strategies for the plaintiff to increase 
            their chances of winning or achieving a favorable settlement.

            Plaintiff's Strategies:
            {plaintiff_strategies}

            Defendant's Strategies:
            {defendant_strategies}

            Please provide a detailed game theory analysis and suggest optimal negotiation strategies for the plaintiff.
            """
            game_theory_prompt = PromptTemplate(template=game_theory_template, input_variables=["plaintiff_strategies", "defendant_strategies"])
            game_theory_chain = LLMChain(llm=llm, prompt=game_theory_prompt)
            
            with st.spinner('Generating game theory analysis...'):
                try:
                    game_theory_analysis = get_response(game_theory_chain, {
                        "plaintiff_strategies": st.session_state.strategies["plaintiff"],
                        "defendant_strategies": st.session_state.strategies["defendant"]
                    })
                    st.session_state.strategies["game_theorist"] = game_theory_analysis
                    st.session_state.history.append(f"Game Theory Analysis: {game_theory_analysis}")
                    st.write(game_theory_analysis)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

    # Download button for strategies
    if any(st.session_state.strategies.values()):
        strategies_text = "\n\n".join([f"{key.capitalize()} Strategies:\n{value}" for key, value in st.session_state.strategies.items() if value])
        bytes_io = io.BytesIO(strategies_text.encode())
        st.download_button(
            label="Download Strategies",
            data=bytes_io,
            file_name="legal_strategies.txt",
            mime="text/plain"
        )

# Additional question section (using document store)
additional_question = st.text_input("Ask additional questions:")
if additional_question:
    if st.button("Get Answer"):
        with st.spinner('Generating answer...'):
            try:
                answer = get_answer_from_docs(additional_question)
                st.session_state.history.append(f"Q: {additional_question}\nA: {answer}")
                st.write(answer)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
